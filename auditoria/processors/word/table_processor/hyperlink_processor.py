"""
Procesador de hipervínculos para documentos Word.
Contiene funciones para aplicar hipervínculos automáticos basados en nomenclatura.
"""

import re
import logging
from django.conf import settings
from docx.oxml.shared import qn, OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

logger = logging.getLogger(__name__)

def apply_hyperlinks_to_document(doc, audit_id, document_name, document_path, nomenclature_config):
    """
    Aplica hipervínculos a las referencias en el documento según la nomenclatura correspondiente
    al nombre del archivo.
    
    Args:
        doc: Documento Word
        audit_id: ID de la auditoría
        document_name: Nombre del documento (sin ruta)
        document_path: Ruta completa del documento
        nomenclature_config: Configuración de nomenclatura obtenida externamente
    """
    if not nomenclature_config:
        return doc
        
    prefix = nomenclature_config.get("prefix")
    max_range = nomenclature_config.get("max_range", 20)
    
    if not prefix:
        return doc
        
    # Recorrer las referencias y aplicar hipervínculos
    hyperlinks_added = 0
    
    # Procesar texto en párrafos
    for paragraph in doc.paragraphs:
        hyperlinks_added += process_paragraph_for_hyperlinks(paragraph, prefix, max_range, audit_id)
    
    # Procesar texto en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    hyperlinks_added += process_paragraph_for_hyperlinks(paragraph, prefix, max_range, audit_id)
    
    return doc

def process_paragraph_for_hyperlinks(paragraph, prefix, max_range, audit_id):
    """
    Procesa un párrafo buscando referencias que coincidan con el patrón y les aplica hipervínculos.
    Se enfoca exclusivamente en el formato estándar (A-1, R-10, etc.).
    
    Antes de procesar, unifica los runs para evitar problemas con patrones divididos entre runs.
    
    Returns:
        int: Número de hipervínculos añadidos
    """
    # Verificar si hay patrones en el texto antes de procesarlo
    potential_patterns = [f"{prefix}-{i}" for i in range(1, max_range + 1)]
    if not any(pattern in paragraph.text for pattern in potential_patterns):
        return 0  # No hay patrones, salir rápidamente
    
    # AQUÍ ESTÁ LA SOLUCIÓN: Unificar todos los runs del párrafo en uno solo
    # Para evitar problemas con patrones divididos en múltiples runs
    
    # 1. Guardar el texto completo
    text = paragraph.text
    
    # 2. Guardar propiedades de estilo del primer run si existe
    style_props = {}
    if paragraph.runs:
        first_run = paragraph.runs[0]
        style_props = {
            'bold': first_run.bold,
            'italic': first_run.italic,
            'underline': first_run.underline,
            'font_size': first_run.font.size,
            'font_name': first_run.font.name,
            'font_color': first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
        }
    
    # 3. Limpiar todos los runs existentes
    for i in range(len(paragraph.runs)-1, -1, -1):
        paragraph._p.remove(paragraph.runs[i]._r)
    
    # 4. Crear un solo run con todo el texto
    new_run = paragraph.add_run(text)
    
    # 5. Aplicar propiedades de estilo guardadas
    if style_props:
        new_run.bold = style_props['bold']
        new_run.italic = style_props['italic']
        new_run.underline = style_props['underline']
        new_run.font.size = style_props['font_size']
        new_run.font.name = style_props['font_name']
        if style_props['font_color']:
            new_run.font.color.rgb = style_props['font_color']
    
    hyperlinks_added = 0
    
    # Un patrón ya procesado no debe procesarse nuevamente
    processed_patterns = set()
    
    # Formato estándar (PREFIX-NUMBER)
    for i in range(1, max_range + 1):
        pattern = f"{prefix}-{i}"
        
        # Usamos regex para encontrar coincidencias exactas de palabras
        # \b representa límite de palabra para evitar coincidencias parciales
        regex_pattern = r'\b' + re.escape(pattern) + r'\b'
        
        # Buscar coincidencias en el texto
        matches = list(re.finditer(regex_pattern, paragraph.text))
        
        # Procesar cada coincidencia encontrada
        if matches:
            for match in matches:
                hyperlinks_added = process_matched_hyperlink(paragraph, match, pattern, processed_patterns, hyperlinks_added, audit_id)
    
    return hyperlinks_added

def process_matched_hyperlink(paragraph, match, pattern, processed_patterns, hyperlinks_added, audit_id):
    """
    Procesa una coincidencia encontrada y aplica el hipervínculo.
    Función helper para evitar duplicación de código.
    """
    # No procesar patrones ya procesados
    if pattern in processed_patterns:
        return hyperlinks_added
        
    # Añadir información de depuración
    
    # Información detallada sobre la coincidencia
    start_idx = match.start()
    end_idx = match.end()
    matched_text = paragraph.text[start_idx:end_idx]
    
    # Para los patrones A-16 a A-18, añadir información adicional
    
    processed_patterns.add(pattern)
    
    # Encontrar el run que contiene el patrón
    current_idx = 0
    target_run = None
    target_run_start = 0
    
    for j, run in enumerate(paragraph.runs):
        run_length = len(run.text)
        if current_idx <= start_idx < current_idx + run_length:
            target_run = run
            target_run_start = current_idx
            break
        current_idx += run_length
    
    if target_run:
        # Crear la URL
        url = f"{settings.BASE_URL}/auditoria/download/{audit_id}/{pattern}"
        
        # Extraer el texto que coincide con el patrón
        run_text = target_run.text
        pattern_start_in_run = start_idx - target_run_start
        pattern_end_in_run = min(end_idx - target_run_start, len(run_text))
        
        # Si el patrón está completamente dentro del run
        if pattern_start_in_run >= 0 and pattern_end_in_run <= len(run_text):
            # Guardar el texto y estilo del run
            text_before = run_text[:pattern_start_in_run]
            text_pattern = run_text[pattern_start_in_run:pattern_end_in_run]
            text_after = run_text[pattern_end_in_run:]
            
            # Propiedades de estilo del run original
            original_bold = target_run.bold
            original_italic = target_run.italic
            original_underline = target_run.underline
            original_font_size = target_run.font.size
            original_font_name = target_run.font.name
            original_font_color = target_run.font.color.rgb if target_run.font.color and target_run.font.color.rgb else None
            
            # Eliminar el run original
            paragraph._p.remove(target_run._r)
            
            # Crear nuevos runs: antes del patrón, patrón con hipervínculo y después del patrón
            if text_before:
                run_before = paragraph.add_run(text_before)
                # Aplicar estilo del run original
                run_before.bold = original_bold
                run_before.italic = original_italic
                run_before.underline = original_underline
                run_before.font.size = original_font_size
                run_before.font.name = original_font_name
                if original_font_color:
                    run_before.font.color.rgb = original_font_color
            
            # --- MÉTODO MEJORADO PARA AGREGAR HIPERVÍNCULOS ---
            # Crear el run para el hipervínculo
            hyperlink_run = paragraph.add_run(text_pattern)
            hyperlink_run.bold = original_bold
            hyperlink_run.italic = original_italic
            hyperlink_run.underline = True  # Subrayar el hipervínculo
            hyperlink_run.font.size = original_font_size
            hyperlink_run.font.name = original_font_name
            if original_font_color:
                hyperlink_run.font.color.rgb = original_font_color
            
            # Crear el hipervínculo manualmente usando la API XML
            r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
            
            # Crear el elemento de hipervínculo
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            hyperlink.set(qn('w:history'), '1')  # Para marcar como visitado
            
            # Mover el run del texto dentro del hipervínculo
            hyperlink.append(hyperlink_run._r)
            
            # Reemplazar el run por el hipervínculo
            paragraph._p.append(hyperlink)
            
            # Remover el run original del párrafo ya que ahora está dentro del hipervínculo
            if hyperlink_run._r in paragraph._p:
                paragraph._p.remove(hyperlink_run._r)
            
            if text_after:
                run_after = paragraph.add_run(text_after)
                run_after.bold = original_bold
                run_after.italic = original_italic
                run_after.underline = original_underline
                run_after.font.size = original_font_size
                run_after.font.name = original_font_name
                if original_font_color:
                    run_after.font.color.rgb = original_font_color
            
            hyperlinks_added += 1
            
    return hyperlinks_added
