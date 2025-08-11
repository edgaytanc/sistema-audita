import logging
import re
from docx.text.paragraph import Paragraph
from ..shared.text_replacer import replace_text
import os
from django.conf import settings

logger = logging.getLogger(__name__)

def set_text_with_style_from_reference_cell(target_cell, text, reference_cell):
    """
    Inserta texto en `target_cell` con el mismo estilo del primer run de `reference_cell`.
    """
    target_cell.text = ""  # Limpiar contenido anterior

    ref_paragraphs = reference_cell.paragraphs
    if ref_paragraphs and ref_paragraphs[0].runs:
        ref_run = ref_paragraphs[0].runs[0]
        paragraph = target_cell.paragraphs[0]
        new_run = paragraph.add_run(text)
        # Copiar estilos básicos
        new_run.bold = ref_run.bold
        new_run.italic = ref_run.italic
        new_run.underline = ref_run.underline
        new_run.font.size = ref_run.font.size
        new_run.font.name = ref_run.font.name
        new_run.font.color.rgb = ref_run.font.color.rgb
    else:
        # Si no se encuentra un estilo de referencia válido, insertar texto plano
        target_cell.text = text

def replace_text_preserving_format(cell, old_text, new_text):
    """
    Reemplaza texto en una celda preservando todos los elementos de formato.
    
    En lugar de asignar directamente text a la celda (lo que destruye el formato),
    esta función trabaja con párrafos y runs individuales.
    """
    if not old_text or old_text == new_text:
        return False

    changed = False
    # Trabajamos a nivel de párrafos para preservar la estructura
    for paragraph in cell.paragraphs:
        if paragraph.text and old_text in paragraph.text:
            # Almacenamos información sobre los runs originales
            runs_info = []
            for run in paragraph.runs:
                runs_info.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'size': run.font.size,
                    'name': run.font.name,
                    'color': run.font.color.rgb
                })
            
            # Reemplazamos el texto completo del párrafo
            new_paragraph_text = paragraph.text.replace(old_text, new_text)
            
            # Limpiamos el párrafo
            for i in range(len(paragraph.runs)-1, -1, -1):
                paragraph._p.remove(paragraph.runs[i]._r)
            
            # Si el párrafo tiene una estructura simple, simplemente añadimos un run
            if len(runs_info) == 1:
                run = paragraph.add_run(new_paragraph_text)
                run.bold = runs_info[0]['bold']
                run.italic = runs_info[0]['italic']
                run.underline = runs_info[0]['underline']
                run.font.size = runs_info[0]['size']
                run.font.name = runs_info[0]['name']
                run.font.color.rgb = runs_info[0]['color']
            # Si es más complejo, reconstruimos los runs manteniendo el formato
            else:
                paragraph.add_run(new_paragraph_text)
            
            changed = True
    
    return changed

def process_tables(doc, replacements, tables_config=None, document_name=None, document_path=None, audit_id=None):
    """
    Procesa y reemplaza texto en tablas de un documento Word.
    
    Args:
        doc: Documento Word a procesar
        replacements: Diccionario con los valores a reemplazar
        tables_config: Configuración específica para tablas (opcional)
        document_name: Nombre del documento (para hipervínculos)
        document_path: Ruta del documento (para hipervínculos)
        audit_id: ID de la auditoría (para hipervínculos)
        
    Returns:
        Documento Word procesado
    """
    patrones = tables_config.get("patrones", [])
    patrones_regex = tables_config.get("patrones_regex", [])
    processed_cells = {}  # Diccionario para rastrear celdas ya procesadas

    for idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Generar un ID único para esta celda
                cell_id = f"table_{idx}_row_{row_idx}_col_{col_idx}"
                
                # Si la celda ya fue procesada, continuamos con la siguiente
                if cell_id in processed_cells:
                    continue
                    
                cell_text = cell.text.strip().replace('\n', ' ').replace('  ', ' ')

                # Caso 1: reemplazo en celda adyacente con estilo
                for patron in patrones:
                    buscar = patron.get("buscar")
                    clave = patron.get("valor")
                    valor = replacements.get(clave, "")

                    if cell_text == buscar:
                        try:
                            next_cell = row.cells[col_idx + 1]
                            next_cell_id = f"table_{idx}_row_{row_idx}_col_{col_idx + 1}"
                            set_text_with_style_from_reference_cell(next_cell, valor, cell)
                            processed_cells[cell_id] = processed_cells[next_cell_id] = True
                        except IndexError:
                            pass

                # Si la celda ya fue procesada, saltamos los otros casos
                if cell_id in processed_cells:
                    continue

                # Caso 2: reemplazo parcial dentro de la misma celda (directo)
                for patron in patrones:
                    buscar = patron.get("buscar")
                    if buscar in cell_text:
                        nuevo_texto = replace_text(cell_text, replacements)
                        if nuevo_texto != cell_text:
                            replace_text_preserving_format(cell, cell_text, nuevo_texto)
                            processed_cells[cell_id] = True
                
                # Caso 3: aplicar reemplazos regex desde JSON
                if not cell_id in processed_cells and patrones_regex:
                    texto_original = cell.text
                    patrones_aplicados = set()
                    
                    for patron_regex in patrones_regex:
                        pattern = patron_regex.get("pattern")
                        key = patron_regex.get("reemplazar_por")
                        
                        patron_id = f"{pattern}_{key}"
                        if patron_id in patrones_aplicados:
                            continue
                            
                        valor = replacements.get(key, "")
                        
                        if pattern and valor and cell.text:
                            match = re.search(pattern, cell.text)
                            if match:
                                patrones_aplicados.add(patron_id)
                                matched_text = cell.text[match.start():match.end()]
                                replace_text_preserving_format(cell, matched_text, valor)
                                processed_cells[cell_id] = True
    
    # Aplicar hipervínculos si hay información suficiente
    if audit_id and document_name:
        doc = apply_hyperlinks_to_document(doc, audit_id, document_name, document_path)

    return doc

def apply_hyperlinks_to_document(doc, audit_id, document_name, document_path):
    """
    Aplica hipervínculos a las referencias en el documento según la nomenclatura correspondiente
    al nombre del archivo.
    
    Args:
        doc: Documento Word
        audit_id: ID de la auditoría
        document_name: Nombre del documento (sin ruta)
        document_path: Ruta completa del documento
    """
    # Definir la nomenclatura y rango según el nombre del archivo
    nomenclature_config = get_nomenclature_config(document_name, document_path)
    
    if not nomenclature_config:
        return doc
        
    prefix = nomenclature_config.get("prefix")
    max_range = nomenclature_config.get("max_range", 20)
    
    if not prefix:
        return doc
        
    # Recorrer las referencias y aplicar hipervínculos
    hyperlinks_added = 0
    
    # Procesar texto en párrafos
    for paragraph in doc.paragraphs:
        hyperlinks_added += process_paragraph_for_hyperlinks(paragraph, prefix, max_range, audit_id)
    
    # Procesar texto en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    hyperlinks_added += process_paragraph_for_hyperlinks(paragraph, prefix, max_range, audit_id)
    
    return doc

def process_paragraph_for_hyperlinks(paragraph, prefix, max_range, audit_id):
    """
    Procesa un párrafo buscando referencias que coincidan con el patrón y les aplica hipervínculos.
    Se enfoca exclusivamente en el formato estándar (A-1, R-10, etc.).
    
    Antes de procesar, unifica los runs para evitar problemas con patrones divididos entre runs.
    
    Returns:
        int: Número de hipervínculos añadidos
    """
    import re
    from docx.oxml.shared import qn, OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    
    # Verificar si hay patrones en el texto antes de procesarlo
    potential_patterns = [f"{prefix}-{i}" for i in range(1, max_range + 1)]
    if not any(pattern in paragraph.text for pattern in potential_patterns):
        return 0  # No hay patrones, salir rápidamente
    
    # AQUÍ ESTÁ LA SOLUCIÓN: Unificar todos los runs del párrafo en uno solo
    # Para evitar problemas con patrones divididos en múltiples runs
    
    # 1. Guardar el texto completo
    text = paragraph.text
    
    # 2. Guardar propiedades de estilo del primer run si existe
    style_props = {}
    if paragraph.runs:
        first_run = paragraph.runs[0]
        style_props = {
            'bold': first_run.bold,
            'italic': first_run.italic,
            'underline': first_run.underline,
            'font_size': first_run.font.size,
            'font_name': first_run.font.name,
            'font_color': first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
        }
    
    # 3. Limpiar todos los runs existentes
    for i in range(len(paragraph.runs)-1, -1, -1):
        paragraph._p.remove(paragraph.runs[i]._r)
    
    # 4. Crear un solo run con todo el texto
    new_run = paragraph.add_run(text)
    
    # 5. Aplicar propiedades de estilo guardadas
    if style_props:
        new_run.bold = style_props['bold']
        new_run.italic = style_props['italic']
        new_run.underline = style_props['underline']
        new_run.font.size = style_props['font_size']
        new_run.font.name = style_props['font_name']
        if style_props['font_color']:
            new_run.font.color.rgb = style_props['font_color']
    
    hyperlinks_added = 0
    
    # Un patrón ya procesado no debe procesarse nuevamente
    processed_patterns = set()
    
    # Formato estándar (PREFIX-NUMBER)
    for i in range(1, max_range + 1):
        pattern = f"{prefix}-{i}"
        
        # Usamos regex para encontrar coincidencias exactas de palabras
        # \b representa límite de palabra para evitar coincidencias parciales
        regex_pattern = r'\b' + re.escape(pattern) + r'\b'
        
        # Buscar coincidencias en el texto
        matches = list(re.finditer(regex_pattern, paragraph.text))
        
        # Procesar cada coincidencia encontrada
        if matches:
            for match in matches:
                hyperlinks_added = process_matched_hyperlink(paragraph, match, pattern, processed_patterns, hyperlinks_added, audit_id)
    
    return hyperlinks_added

def process_matched_hyperlink(paragraph, match, pattern, processed_patterns, hyperlinks_added, audit_id):
    """
    Procesa una coincidencia encontrada y aplica el hipervínculo.
    Función helper para evitar duplicación de código.
    """
    from docx.oxml.shared import qn, OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    
    # No procesar patrones ya procesados
    if pattern in processed_patterns:
        return hyperlinks_added
        
    # Añadir información de depuración
    
    # Información detallada sobre la coincidencia
    start_idx = match.start()
    end_idx = match.end()
    matched_text = paragraph.text[start_idx:end_idx]
    
    # Para los patrones A-16 a A-18, añadir información adicional
    
    processed_patterns.add(pattern)
    
    # Encontrar el run que contiene el patrón
    current_idx = 0
    target_run = None
    target_run_start = 0
    
    for j, run in enumerate(paragraph.runs):
        run_length = len(run.text)
        if current_idx <= start_idx < current_idx + run_length:
            target_run = run
            target_run_start = current_idx
            break
        current_idx += run_length
    
    if target_run:
        # Crear la URL
        url = f"{settings.BASE_URL}/auditoria/download/{audit_id}/{pattern}"
        
        # Extraer el texto que coincide con el patrón
        run_text = target_run.text
        pattern_start_in_run = start_idx - target_run_start
        pattern_end_in_run = min(end_idx - target_run_start, len(run_text))
        
        # Si el patrón está completamente dentro del run
        if pattern_start_in_run >= 0 and pattern_end_in_run <= len(run_text):
            # Guardar el texto y estilo del run
            text_before = run_text[:pattern_start_in_run]
            text_pattern = run_text[pattern_start_in_run:pattern_end_in_run]
            text_after = run_text[pattern_end_in_run:]
            
            # Propiedades de estilo del run original
            original_bold = target_run.bold
            original_italic = target_run.italic
            original_underline = target_run.underline
            original_font_size = target_run.font.size
            original_font_name = target_run.font.name
            original_font_color = target_run.font.color.rgb if target_run.font.color and target_run.font.color.rgb else None
            
            # Eliminar el run original
            paragraph._p.remove(target_run._r)
            
            # Crear nuevos runs: antes del patrón, patrón con hipervínculo y después del patrón
            if text_before:
                run_before = paragraph.add_run(text_before)
                # Aplicar estilo del run original
                run_before.bold = original_bold
                run_before.italic = original_italic
                run_before.underline = original_underline
                run_before.font.size = original_font_size
                run_before.font.name = original_font_name
                if original_font_color:
                    run_before.font.color.rgb = original_font_color
            
            # --- MÉTODO MEJORADO PARA AGREGAR HIPERVÍNCULOS ---
            # Crear el run para el hipervínculo
            hyperlink_run = paragraph.add_run(text_pattern)
            hyperlink_run.bold = original_bold
            hyperlink_run.italic = original_italic
            hyperlink_run.underline = True  # Subrayar el hipervínculo
            hyperlink_run.font.size = original_font_size
            hyperlink_run.font.name = original_font_name
            if original_font_color:
                hyperlink_run.font.color.rgb = original_font_color
            
            # Crear el hipervínculo manualmente usando la API XML
            r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
            
            # Crear el elemento de hipervínculo
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            hyperlink.set(qn('w:history'), '1')  # Para marcar como visitado
            
            # Mover el run del texto dentro del hipervínculo
            hyperlink.append(hyperlink_run._r)
            
            # Reemplazar el run por el hipervínculo
            paragraph._p.append(hyperlink)
            
            # Remover el run original del párrafo ya que ahora está dentro del hipervínculo
            if hyperlink_run._r in paragraph._p:
                paragraph._p.remove(hyperlink_run._r)
            
            if text_after:
                run_after = paragraph.add_run(text_after)
                run_after.bold = original_bold
                run_after.italic = original_italic
                run_after.underline = original_underline
                run_after.font.size = original_font_size
                run_after.font.name = original_font_name
                if original_font_color:
                    run_after.font.color.rgb = original_font_color
            
            hyperlinks_added += 1
            
    return hyperlinks_added

def get_nomenclature_config(document_name, document_path=None):
    """
    Devuelve la configuración de nomenclatura para un documento en particular.
    
    Args:
        document_name: Nombre del documento (sin ruta)
        document_path: Ruta completa del documento, necesaria para mapeos basados en ruta
    
    Returns:
        dict: Configuración de nomenclatura con prefix y max_range o None si no coincide
    """
    # Primero intentamos con la configuración existente
    document_name_lower = document_name.lower()
    
    # Configuración original para documentos estándar
    nomenclature_map = {
        "1 programa.docx": {"prefix": "A", "max_range": 18},
        "1 programa de auditoria cuentas por cobrar.docx": {"prefix": "B", "max_range": 16},
        "1 programa de auditoria inversiones.docx": {"prefix": "C", "max_range": 6},
        "1 programa de auditoria inventarios.docx": {"prefix": "D", "max_range": 16},
        "1 programa de auditoria construcciones en proceso.docx": {"prefix": "E", "max_range": 13},
        "1 programa de auditoria activos fijos.docx": {"prefix": "F", "max_range": 13},
        "1 programa de activo intangible.docx": {"prefix": "G", "max_range": 10},
        "1 programa de auditoria cuentas por pagar.docx": {"prefix": "H", "max_range": 19},
        "1 programa de auditoria prestamos por pagar.docx": {"prefix": "I", "max_range": 8},
        "1 programa de auditoria patrimonio.docx": {"prefix": "J", "max_range": 8},
        "1 programa de auditoria estado resultados.docx": {"prefix": "R", "max_range": 13},
    }
    
    # Nueva configuración para documentos de auditoría interna
    nomenclature_map_interna = {
        "6 AUDITORIA DE PROCESOS": {
            "1 Operativo": {
                "1 Gestión de Inventarios y Almacenes": {
                    "4 Programa de Auditoría.docx": {
                        "prefix": "B",
                        "max_range": 12  # B-6, B-7, B-8, B-9, B-10
                    }
                }
            },
            "2 Comercial": {
                "4 Programa de Auditoría.docx": {
                    "prefix": "D",
                    "max_range": 12  # D-5, D-6, D-7, D-8, D-9
                }
            },
            "3 Financiero": {
                "1 AUDITORIA PROCESOS TESORERIA": {
                    "4 Programa de Auditoría.docx": {
                        "prefix": "E",
                        "max_range": 12  # E-5, E-6, E-7, E-8, E-9, E-10
                    }
                },
                "2 AUDITORIA PROCESOS CONTABILIDAD": {
                    "4 Programa de Auditoría.docx": {
                        "prefix": "F",
                        "max_range": 12  # F-5, F-6, F-7, F-8, F-9, F-10
                    }
                },
                "3 AUDITORIA PROCESOS COMPRAS": {
                    "4 Programa de Auditoría.docx": {
                        "prefix": "H",
                        "max_range": 12  # H-5, H-6, H-7, H-8, H-9, H-10
                    }
                }
            },
            "4 Recursos Humanos": {
                "4 Programa de Auditoría.docx": {
                    "prefix": "K",
                    "max_range": 12  # K-5, K-6, K-7, K-8, K-9, K-10, K-11
                }
            },
            "5 Tecnología": {
                "4 Programa de Auditoría.docx": {
                    "prefix": "L",
                    "max_range": 12  # L-5, L-6, L-7, L-8, L-9
                }
            },
            "6 Legal": {
                "4 Programa de Auditoría.docx": {
                    "prefix": "M",
                    "max_range": 12  # M-5, M-6, M-7, M-8, M-9
                }
            },
            "7 Servicios Públicos": {
                "4 Programa de Auditoría.docx": {
                    "prefix": "N",
                    "max_range": 12  # N-5, N-6, N-7, N-8, N-9
                }
            },
            "8 Obras Públicas": {
                "4 Programa de Auditoría.docx": {
                    "prefix": "O",
                    "max_range": 12  # O-5, O-6, O-7, O-8, O-9, O-10
                }
            },
            "9 Gestión de Ahorro": {
                "4 Programa de Auditoría.docx": {
                    "prefix": "P",
                    "max_range": 12  # P-5, P-6, P-7, P-8, P-9
                }
            },
            "10 Gestión de Créditos": {
                "4 Programa de Auditoría.docx": {
                    "prefix": "Q",
                    "max_range": 12  # Q-5, Q-6, Q-7, Q-8, Q-9, Q-10, Q-11
                }
            }
        }
    }
    
    # Verificar primero en el mapa estándar
    config = nomenclature_map.get(document_name_lower)
    if config:
        return config
    
    # Si no se encontró y tenemos la ruta del documento, verificamos en la estructura de auditoría interna
    if document_path and "programa de auditor" in document_name.lower():
        # Normalizar la ruta para la búsqueda
        normalized_path = document_path.replace('\\', '/').lower()
        
        # Verificar si la ruta contiene "6 AUDITORIA DE PROCESOS"
        if "6 auditoria de procesos" in normalized_path:
            
            # Mapa directo de patrones de ruta a prefijos (para búsqueda simplificada)
            directos = {
                # Áreas con subáreas específicas 
                "1 operativo/1 gestión de inventarios": "B",
                "3 financiero/1 auditoria procesos tesoreria": "E",
                "3 financiero/2 auditoria procesos contabilidad": "F",
                "3 financiero/3 auditoria procesos compras": "H",
                
                # Áreas sin subáreas (solo el nombre del área)
                "2 comercial": "D",
                "4 recursos humanos": "K",
                "5 tecnología": "L",
                "5 tecnologia": "L",
                "6 legal": "M",
                "7 servicios públicos": "N",
                "7 servicios publicos": "N",
                "8 obras públicas": "O",
                "8 obras publicas": "O",
                "9 gestión de ahorro": "P",
                "9 gestion de ahorro": "P",
                "10 gestión de créditos": "Q",
                "10 gestion de creditos": "Q"
            }
            
            # Primero intentamos la coincidencia directa por patrones de ruta
            for pattern, prefix in directos.items():
                if pattern in normalized_path:
                    max_range = 12  # Valor estándar para todos
                    config = {"prefix": prefix, "max_range": max_range}
                    return config
            
            # Si no encontramos coincidencia directa, buscamos por áreas específicas
            for area_name, area_config in nomenclature_map_interna["6 AUDITORIA DE PROCESOS"].items():
                area_pattern = area_name.lower()
                if area_pattern in normalized_path:
                    
                    # Verificar si el área tiene subáreas o documentos directamente
                    if isinstance(next(iter(area_config.values())), dict):
                        # El área tiene subáreas
                        # Verificar si alguna subárea está en la ruta
                        for subarea_name, subarea_config in area_config.items():
                            if subarea_name.lower() in normalized_path:
                                
                                # Verificar si el documento coincide
                                for doc_name, doc_config in subarea_config.items():
                                    doc_name_clean = doc_name.lower().replace('í', 'i').replace('á', 'a')
                                    doc_search = document_name.lower().replace('í', 'i').replace('á', 'a')
                                    
                                    if doc_name_clean in doc_search:
                                        return doc_config
                    else:
                        # El área tiene documentos directamente (sin subáreas)
                        for doc_name, doc_config in area_config.items():
                            doc_name_clean = doc_name.lower().replace('í', 'i').replace('á', 'a')
                            doc_search = document_name.lower().replace('í', 'i').replace('á', 'a')
                            
                            if doc_name_clean in doc_search:
                                return doc_config
    
    return None
